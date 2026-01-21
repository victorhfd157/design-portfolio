#!/usr/bin/env python3
"""
Script COMPLETO para gerar todos os guias HTML a partir dos arquivos Word
"""

from docx import Document
import os
import re

# Template HTML base
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="pt">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Guia da {session_title} | Geração Futuro</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;700;900&family=Space+Grotesk:wght@300;500;700&display=swap" rel="stylesheet">
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        body {{ background-color: #020617; color: white; font-family: 'Outfit', sans-serif; }}
        .glass-card {{ background: rgba(30, 41, 59, 0.4); backdrop-filter: blur(12px); border: 1px solid rgba(255, 255, 255, 0.08); }}
        .content-section h2 {{ color: #a855f7; margin-top: 2rem; margin-bottom: 1rem; font-size: 1.5rem; font-weight: 700; }}
        .content-section h3 {{ color: #22d3ee; margin-top: 1.5rem; margin-bottom: 0.75rem; font-size: 1.25rem; font-weight: 600; }}
        .content-section p {{ margin-bottom: 1rem; line-height: 1.8; color: #cbd5e1; }}
        .content-section ul, .content-section ol {{ margin-left: 1.5rem; margin-bottom: 1rem; color: #cbd5e1; }}
        .content-section li {{ margin-bottom: 0.5rem; line-height: 1.6; }}
        .activity-box {{ background: rgba(168, 85, 247, 0.1); border-left: 4px solid #a855f7; padding: 1.5rem; margin: 1.5rem 0; border-radius: 0.5rem; }}
    </style>
</head>
<body class="antialiased">
    <div class="fixed inset-0 -z-10">
        <div class="absolute inset-0 bg-gradient-to-br from-purple-950/20 via-slate-950 to-cyan-950/20"></div>
        <div class="absolute inset-0" style="background-image: radial-gradient(1px 1px at 50px 50px, rgba(255,255,255,0.3) 50%, transparent 100%); background-size: 500px 500px; opacity: 0.3;"></div>
    </div>

    <nav class="sticky top-0 z-50 bg-black/50 backdrop-blur-xl border-b border-white/5 py-4 px-6">
        <div class="max-w-5xl mx-auto flex items-center justify-between">
            <a href="../../index.html" class="flex items-center gap-3 text-white hover:text-purple-300 transition-colors">
                <svg class="w-6 h-6" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M10 19l-7-7m0 0l7-7m-7 7h18"/>
                </svg>
                <span class="font-medium">Voltar ao Hub</span>
            </a>
            <div class="flex items-center gap-3">
                <span class="text-xs uppercase tracking-widest text-cyan-400 font-bold">{module_title}</span>
                <a href="../../modulo{module_num}/sessao{session_num}/" class="px-4 py-2 rounded-full bg-purple-600 hover:bg-purple-700 text-white text-sm font-medium transition-all">
                    Ver Apresentação
                </a>
            </div>
        </div>
    </nav>

    <main class="max-w-5xl mx-auto px-6 py-12">
        <div class="mb-12">
            <div class="flex items-center gap-3 mb-4">
                <span class="px-3 py-1 rounded-full bg-purple-500/20 text-purple-300 text-xs font-bold uppercase tracking-wider">Guia do Formador</span>
            </div>
            <h1 class="text-5xl font-display font-black text-white mb-4">{session_title}</h1>
            <p class="text-xl text-slate-400">Guia completo para conduzir esta sessão</p>
        </div>

        <div class="glass-card rounded-2xl p-8 content-section">
            {content}
        </div>

        <div class="mt-8 glass-card rounded-2xl p-6 flex items-center justify-between">
            <div>
                <h3 class="text-white font-bold mb-1">Documento Original</h3>
                <p class="text-slate-400 text-sm">Download do ficheiro Word completo</p>
            </div>
            <a href="../modulo{module_num}/{word_filename}" download class="px-6 py-3 rounded-full bg-white/10 hover:bg-white/20 border border-white/20 hover:border-white/30 transition-all flex items-center gap-2 text-white font-medium">
                <svg class="w-5 h-5" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M12 10v6m0 0l-3-3m3 3l3-3m2 8H7a2 2 0 01-2-2V5a2 2 0 012-2h5.586a1 1 0 01.707.293l5.414 5.414a1 1 0 01.293.707V19a2 2 0 01-2 2z"/>
                </svg>
                Download Word
            </a>
        </div>
    </main>

    <footer class="mt-20 border-t border-white/5 py-8">
        <div class="max-w-5xl mx-auto px-6 text-center text-slate-500 text-sm">
            <p>act.academy | Geração Futuro © 2026</p>
        </div>
    </footer>
</body>
</html>"""

def process_word_to_html(docx_path):
    """Extrai conteúdo do Word e converte para HTML"""
    doc = Document(docx_path)
    html_parts = []
    current_list = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or text == "":
            continue
        
        # Detecta títulos (negrito ou Heading)
        is_heading = para.style.name.startswith('Heading') or (para.runs and any(run.bold and len(run.text.strip()) > 5 for run in para.runs))
        
        if is_heading:
            # Fecha lista se aberta
            if current_list:
                html_parts.append(f"</{current_list}>")
                current_list = None
            
            # Adiciona ícone baseado no conteúdo
            icon = "📝"
            if any(kw in text.lower() for kw in ['objetivo', 'competência']): icon = "🎯"
            elif any(kw in text.lower() for kw in ['material', 'recurso']): icon = "📦"
            elif any(kw in text.lower() for kw in ['preparação', 'antes']): icon = "🔧"
            elif any(kw in text.lower() for kw in ['atividade', 'exercício', 'dinâmica']): icon = "🚀"
            elif any(kw in text.lower() for kw in ['avaliação']): icon = "📊"
            
            level = 'h2' if 'Sessão' not in text else 'h3'
            html_parts.append(f"<{level}>{icon} {text}</{level}>")
        
        # Detecta listas
        elif text.startswith('•') or text.startswith('-') or re.match(r'^\d+[\.)]\s', text):
            list_type = 'ol' if re.match(r'^\d+[\.)]\s', text) else 'ul'
            
            if current_list != list_type:
                if current_list:
                    html_parts.append(f"</{current_list}>")
                html_parts.append(f"<{list_type} class='list-disc'>")
                current_list = list_type
            
            clean_text = re.sub(r'^[•\-\d+\.)]\s*', '', text)
            html_parts.append(f"<li>{clean_text}</li>")
        
        else:
            # Fecha lista se aberta
            if current_list:
                html_parts.append(f"</{current_list}>")
                current_list = None
            
            # Parágrafo normal
            html_parts.append(f"<p>{text}</p>")
    
    # Fecha lista final se necessário
    if current_list:
        html_parts.append(f"</{current_list}>")
    
    return "\n".join(html_parts)

# Mapeia arquivos Word para módulos/sessões
word_files = {
    1: [
        ("Sessão 1.docx", 1),
        ("Sessão 2.docx", 2),
        ("Sessão 3.docx", 3),
    ],
    2: [
        ("M2 - Sessão 1.docx", 1),
        ("M2 - Sessão 2.docx", 2),
        ("M2 - Sessão 3.docx", 3),
        ("M2 - Sessão 4.docx", 4),
        ("M2 - Sessão 5.docx", 5),
    ],
    3: [
        ("M3 - Sessão 1.docx", 1),
        ("M3 - Sessão 2.docx", 2),
        ("M3 - Sessão 3.docx", 3),
        ("M3 - Sessão 4.docx", 4),
        ("M3 - Sessão 5.docx", 5),
        ("M3 - Sessão 6.docx", 6),
        ("M3 - Sessão 7.docx", 7),
        ("M3 - Sessão 8.docx", 8),
    ],
    4: [
        ("M4 - Sessão 1.docx", 1),
        ("M4 - Sessão 2.docx", 2),
        ("M4 - Sessão 3.docx", 3),
        ("M4 - Sessão 4.docx", 4),
        ("M4 - Sessão 5.docx", 5),
        ("M4 - Sessão 6.docx", 6),
        ("M4 - Sessão 7.docx", 7),
    ],
    5: [
        ("M5 - Sessão 1.docx", 1),
        ("M5 - Sessão 2.docx", 2),
        ("M5 - Sessão 3.docx", 3),
        ("M5 - Sessão 4.docx", 4),
        ("M5 - Sessão 6.docx", 6),
    ]
}

# Processa todos os arquivos
total_created = 0
total_errors = 0

for module_num, sessions in word_files.items():
    print(f"\n📚 Processando Módulo {module_num}...")
    
    for word_filename, session_num in sessions:
        docx_path = f"resources/modulo{module_num}/{word_filename}"
        output_path = f"resources/modulo{module_num}/sessao{session_num}-guia.html"
        
        if not os.path.exists(docx_path):
            print(f"  ⚠️  Arquivo não encontrado: {word_filename}")
            total_errors += 1
            continue
        
        try:
            # Extrai conteúdo
            content_html = process_word_to_html(docx_path)
            
            # Extrai título
            doc = Document(docx_path)
            session_title = f"Sessão {session_num}"
            for para in doc.paragraphs[:10]:
                if "Sessão" in para.text:
                    session_title = para.text.strip().replace("📘", "").replace("Atividade Assíncrona", "").strip()
                    break
            
            # Gera HTML final
            html = HTML_TEMPLATE.format(
                session_title=session_title,
                module_title=f"Módulo {module_num}",
                module_num=module_num,
                session_num=session_num,
                content=content_html,
                word_filename=word_filename
            )
            
            # Salva arquivo
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            print(f"  ✅ {session_title}")
            total_created += 1
            
        except Exception as e:
            print(f"  ❌ Erro em {word_filename}: {e}")
            total_errors += 1

print(f"\n\n🎉 Concluído!")
print(f"   ✅ {total_created} guias criados com sucesso")
print(f"   ❌ {total_errors} erros")
print(f"\n💡 Acesse http://localhost:3000 e teste os guias!")
