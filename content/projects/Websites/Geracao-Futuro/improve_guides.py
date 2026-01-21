#!/usr/bin/env python3
"""
Script MELHORADO para gerar guias HTML com visual aprimorado
"""

from docx import Document
import os
import re

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="pt">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Guia da {session_title} | Geração Futuro</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;700;900&family=Space+Grotesk:wght@300;500;700&display=swap" rel="stylesheet">
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        body {{ background-color: #020617; color: white; font-family: 'Outfit', sans-serif; }
        .glass-card {{ background: rgba(30, 41, 59, 0.6); backdrop-filter: blur(16px); border: 1px solid rgba(255, 255, 255, 0.1); box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3); }
        .content-section h2 {{ color: #a855f7; margin-top: 3rem; margin-bottom: 1.5rem; font-size: 1.75rem; font-weight: 800; border-bottom: 2px solid rgba(168, 85, 247, 0.3); padding-bottom: 0.5rem; }
        .content-section h3 {{ color: #22d3ee; margin-top: 2rem; margin-bottom: 1rem; font-size: 1.35rem; font-weight: 700; }
        .content-section h4 {{ color: #cbd5e1; margin-top: 1.5rem; margin-bottom: 0.75rem; font-size: 1.1rem; font-weight: 600; }
        .content-section p {{ margin-bottom: 1.25rem; line-height: 1.9; color: #cbd5e1; font-size: 1.05rem; }
        .content-section ul, .content-section ol {{ margin-left: 2rem; margin-bottom: 1.5rem; color: #cbd5e1; }
        .content-section li {{ margin-bottom: 0.75rem; line-height: 1.7; }
        .content-section li::marker {{ color: #a855f7; }
        .activity-box {{ background: linear-gradient(135deg, rgba(168, 85, 247, 0.15) 0%, rgba(34, 211, 238, 0.1) 100%); border-left: 4px solid #a855f7; padding: 2rem; margin: 2rem 0; border-radius: 1rem; box-shadow: 0 4px 16px rgba(168, 85, 247, 0.2); }
        .info-box {{ background: rgba(34, 211, 238, 0.1); border-left: 4px solid #22d3ee; padding: 1.5rem; margin: 1.5rem 0; border-radius: 0.75rem; }
        .tip-box {{ background: rgba(251, 191, 36, 0.1); border-left: 4px solid #fbbf24; padding: 1.5rem; margin: 1.5rem 0; border-radius: 0.75rem; }
        .toc {{ position: sticky; top: 100px; max-height: calc(100vh - 120px); overflow-y: auto; }
        .toc::-webkit-scrollbar {{ width: 4px; }
        .toc::-webkit-scrollbar-thumb {{ background: #a855f7; border-radius: 2px; }
        .section-divider {{ height: 1px; background: linear-gradient(90deg, transparent, rgba(168, 85, 247, 0.5), transparent); margin: 3rem 0; }
    </style>
</head>
<body class="antialiased">
    <div class="fixed inset-0 -z-10">
        <div class="absolute inset-0 bg-gradient-to-br from-purple-950/20 via-slate-950 to-cyan-950/20"></div>
        <div class="absolute inset-0" style="background-image: radial-gradient(1px 1px at 50px 50px, rgba(255,255,255,0.3) 50%, transparent 100%); background-size: 500px 500px; opacity: 0.3;"></div>
    </div>

    <nav class="sticky top-0 z-50 bg-black/60 backdrop-blur-xl border-b border-white/10 py-4 px-6 shadow-lg">
        <div class="max-w-7xl mx-auto flex items-center justify-between">
            <a href="../../index.html" class="flex items-center gap-3 text-white hover:text-purple-300 transition-all hover:scale-105">
                <svg class="w-6 h-6" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M10 19l-7-7m0 0l7-7m-7 7h18"/>
                </svg>
                <span class="font-medium">Voltar ao Hub</span>
            </a>
            <div class="flex items-center gap-3">
                <span class="text-xs uppercase tracking-widest text-cyan-400 font-bold">{module_title}</span>
                <a href="../../modulo{module_num}/sessao{session_num}/" class="px-4 py-2 rounded-full bg-gradient-to-r from-purple-600 to-cyan-600 hover:from-purple-700 hover:to-cyan-700 text-white text-sm font-medium transition-all hover:scale-105 shadow-lg">
                    Ver Apresentação
                </a>
            </div>
        </div>
    </nav>

    <main class="max-w-7xl mx-auto px-6 py-12">
        <div class="mb-12">
            <div class="flex items-center gap-3 mb-4">
                <span class="px-4 py-2 rounded-full bg-gradient-to-r from-purple-500/30 to-cyan-500/30 text-purple-200 text-xs font-bold uppercase tracking-wider border border-purple-500/30">Guia do Formador</span>
            </div>
            <h1 class="text-6xl font-display font-black text-white mb-4 bg-gradient-to-r from-purple-400 to-cyan-400 bg-clip-text text-transparent">{session_title}</h1>
            <p class="text-xl text-slate-300">Guia completo para conduzir esta sessão com sucesso</p>
        </div>

        <div class="grid grid-cols-1 lg:grid-cols-4 gap-8">
            <!-- Índice lateral -->
            <aside class="lg:col-span-1">
                <div class="glass-card rounded-2xl p-6 toc">
                    <h3 class="text-white font-bold text-lg mb-4 flex items-center gap-2">
                        <svg class="w-5 h-5 text-cyan-400" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M4 6h16M4 12h16M4 18h16"/>
                        </svg>
                        Navegação
                    </h3>
                    {toc_html}
                </div>
            </aside>

            <!-- Conteúdo principal -->
            <div class="lg:col-span-3">
                <div class="glass-card rounded-2xl p-8 lg:p-12 content-section">
                    {content}
                </div>

                <div class="mt-8 glass-card rounded-2xl p-6 flex flex-col md:flex-row items-center justify-between gap-4">
                    <div>
                        <h3 class="text-white font-bold mb-1">Documento Original</h3>
                        <p class="text-slate-400 text-sm">Faça download do ficheiro Word completo</p>
                    </div>
                    <a href="../modulo{module_num}/{word_filename}" download class="px-6 py-3 rounded-full bg-gradient-to-r from-white/10 to-white/5 hover:from-white/20 hover:to-white/10 border border-white/20 hover:border-white/30 transition-all flex items-center gap-2 text-white font-medium shadow-lg hover:scale-105">
                        <svg class="w-5 h-5" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M12 10v6m0 0l-3-3m3 3l3-3m2 8H7a2 2 0 01-2-2V5a2 2 0 012-2h5.586a1 1 0 01.707.293l5.414 5.414a1 1 0 01.293.707V19a2 2 0 01-2 2z"/>
                        </svg>
                        Download Word
                    </a>
                </div>
            </div>
        </div>
    </main>

    <footer class="mt-20 border-t border-white/10 py-8 bg-black/30">
        <div class="max-w-7xl mx-auto px-6 text-center text-slate-400 text-sm">
            <p>act.academy | Geração Futuro © 2026</p>
        </div>
    </footer>
</body>
</html>"""

def process_word_improved(docx_path):
    """Processa Word com melhor estruturação"""
    doc = Document(docx_path)
    sections = []
    current_section = None
    toc_items = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Detecta seções principais (h2)
        if any(keyword in text for keyword in ['Guião do Formador', 'Objetivos', 'Estrutura Detalhada', 'CONTEÚDO DOS SLIDES', 'FICHA DE ATIVIDADE', 'Teasing', 'Enriquecimento', 'Recursos adicionais']):
            if current_section:
                sections.append(current_section)
            
            section_id = text.lower().replace(' ', '-').replace(':', '')[:30]
            icon = get_icon(text)
            current_section = {
                'id': section_id,
                'title': text,
                'icon': icon,
                'content': [],
                'type': 'section'
            }
            toc_items.append(f'<a href="#{section_id}" class="block py-2 px-3 rounded text-sm text-slate-300 hover:text-white hover:bg-white/10 transition-colors">{icon} {text[:40]}</a>')
        
        elif current_section:
            current_section['content'].append(text)
    
    if current_section:
        sections.append(current_section)
    
    # Gera HTML
    html_parts = []
    for section in sections:
        html_parts.append(f'<h2 id="{section["id"]}">{section["icon"]} {section["title"]}</h2>')
        
        # Processa conteúdo da seção
        in_list = False
        for item in section['content']:
            if item.startswith('•') or item.startswith('-') or re.match(r'^\d+[\.)]\s', item):
                if not in_list:
                    html_parts.append('<ul class="list-disc space-y-2">')
                    in_list = True
                clean = re.sub(r'^[•\-\d+\.)]\s*', '', item)
                html_parts.append(f'<li>{clean}</li>')
            else:
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                
                # Detecta se é subtítulo ou parágrafo
                if len(item) < 80 and (':' in item or item.isupper() or any(emoji in item for emoji in ['🎯', '📝', '🚀', '💡', '🧩'])):
                    html_parts.append(f'<h3>{item}</h3>')
                else:
                    html_parts.append(f'<p>{item}</p>')
        
        if in_list:
            html_parts.append('</ul>')
        
        html_parts.append('<div class="section-divider"></div>')
    
    toc_html = '\n'.join(toc_items) if toc_items else '<p class="text-slate-400 text-sm">Sem índice</p>'
    
    return '\n'.join(html_parts), toc_html

def get_icon(text):
    """Retorna ícone baseado no conteúdo"""
    text_lower = text.lower()
    if 'objetivo' in text_lower: return '🎯'
    if 'material' in text_lower or 'recurso' in text_lower: return '📦'
    if 'preparação' in text_lower or 'guião' in text_lower: return '🔧'
    if 'atividade' in text_lower or 'exercício' in text_lower or 'ficha' in text_lower: return '🚀'
    if 'avaliação' in text_lower: return '📊'
    if 'slide' in text_lower: return '🎬'
    if 'teasing' in text_lower: return '🎮'
    if 'enriquecimento' in text_lower: return '🌟'
    return '📝'

# Processa TODOS os módulos
word_files = {
    1: [("Sessão 1.docx", 1), ("Sessão 2.docx", 2), ("Sessão 3.docx", 3)],
    2: [("M2 - Sessão 1.docx", 1), ("M2 - Sessão 2.docx", 2), ("M2 - Sessão 3.docx", 3), ("M2 - Sessão 4.docx", 4), ("M2 - Sessão 5.docx", 5)],
    3: [("M3 - Sessão 1.docx", 1), ("M3 - Sessão 2.docx", 2), ("M3 - Sessão 3.docx", 3), ("M3 - Sessão 4.docx", 4), ("M3 - Sessão 5.docx", 5), ("M3 - Sessão 6.docx", 6), ("M3 - Sessão 7.docx", 7), ("M3 - Sessão 8.docx", 8)],
    4: [("M4 - Sessão 1.docx", 1), ("M4 - Sessão 2.docx", 2), ("M4 - Sessão 3.docx", 3), ("M4 - Sessão 4.docx", 4), ("M4 - Sessão 5.docx", 5), ("M4 - Sessão 6.docx", 6), ("M4 - Sessão 7.docx", 7)],
    5: [("M5 - Sessão 1.docx", 1), ("M5 - Sessão 2.docx", 2), ("M5 - Sessão 3.docx", 3), ("M5 - Sessão 4.docx", 4), ("M5 - Sessão 6.docx", 6)]
}

total = 0
for module_num, sessions in word_files.items():
    print(f"\n🎨 Melhorando Módulo {module_num}...")
    
    for word_filename, session_num in sessions:
        docx_path = f"resources/modulo{module_num}/{word_filename}"
        output_path = f"resources/modulo{module_num}/sessao{session_num}-guia.html"
        
        if not os.path.exists(docx_path):
            continue
        
        try:
            content_html, toc_html = process_word_improved(docx_path)
            
            doc = Document(docx_path)
            session_title = f"Sessão {session_num}"
            for para in doc.paragraphs[:10]:
                if "Sessão" in para.text:
                    session_title = para.text.strip().replace("📘", "").replace("Atividade Assíncrona", "").strip()
                    break
            
            html = HTML_TEMPLATE.format(
                session_title=session_title,
                module_title=f"Módulo {module_num}",
                module_num=module_num,
                session_num=session_num,
                content=content_html,
                toc_html=toc_html,
                word_filename=word_filename
            )
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            print(f"  ✨ {session_title}")
            total += 1
            
        except Exception as e:
            print(f"  ❌ Erro: {e}")

print(f"\n\n🎉 {total} guias melhorados com sucesso!")
print("💎 Visual aprimorado com:")
print("   • Índice de navegação lateral")
print("   • Tipografia melhorada")
print("   • Cards e boxes destacados")
print("   • Gradientes e sombras")
print("   • Espaçamento otimizado")
